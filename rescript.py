# necessary imports
from docx import Document
from pathlib import Path
import os
import sys

# The function below bolds the passed in paragraph
def bold_paragraph(paragraph):
    bold_para = paragraph.text
    paragraph = paragraph.clear()
    paragraph.add_run(bold_para).bold = True

# The function below deletes the passed in paragraph
def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

# The function below removes the interviewer and participant code paragraphs
# and bolds interviewer text
def clean_transcript(input_path):
    # open doc (object) w path relative to script
    try:
        input_doc = Document(input_path)
    except:
        print("Input doc not found, please place in input folder as .docx file")
        sys.exit(1)

    section = input_doc.sections[0]
    header = section.header

    # interview_content denotes if the curr paragraph is interview content
    # if "Interviewer (" is seen (line 31)
    interview_content = False

    # MS-01-AB
    # iterate through paragraphs in doc
    # https://python-docx.readthedocs.io/en/latest/api/text.html#paragraph-objects
    for para in input_doc.paragraphs:
        # cases bold the interview content, delete the interviewer line,
        # or delete the participant code line
        if "MS-" in para.text:
            interview_content = False
            for run in para.runs:
                run.font.bold = False
        elif interview_content:
            bold_paragraph(para)
            interview_content = False
        elif "Interviewer (" in para.text:
            #delete_paragraph(para)
            interview_content = True

    # save input doc (you may also specify the same input path to overwrite) (line 40)
    input_doc.save("output/" + input_path.split("/")[-1])
    # input_doc.save(input_path)

def main():
    # set input_path (in this case, the file is in the input folder)
    source_dir = Path('input/')
    files = source_dir.iterdir()
    files = source_dir.glob('*.docx')

    for file_name in files:
        f = str(file_name)
        clean_transcript(f)

# https://realpython.com/python-main-function/ -> this runs when you execute "python rescript.py"
if __name__ == "__main__":
    main()








